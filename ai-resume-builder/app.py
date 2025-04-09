from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import pdfkit
import logging
import json
import os
import uuid
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Create the uploads directory if it doesn't exist
if not os.path.exists('uploads'):
    os.makedirs('uploads')

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Hugging Face API configuration
HF_API_URL = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2"
HF_API_TOKEN = os.getenv("HF_API_TOKEN")
if not HF_API_TOKEN:
    raise ValueError("HF_API_TOKEN environment variable is not set")

def generate_text(prompt):
    """Generate enhanced text using Hugging Face's Inference API."""
    try:
        logger.info(f"Sending request to Hugging Face API with prompt length: {len(prompt)}")
        headers = {
            "Authorization": f"Bearer {HF_API_TOKEN}",
            "Content-Type": "application/json"
        }
        
        # Format the prompt for better results
        formatted_prompt = f"""<s>[INST] You are a professional resume writer. Your task is to enhance the following text while maintaining accuracy and professionalism. Keep the core information unchanged but make it more impactful:

{prompt}

Respond with only the enhanced text, no explanations or additional formatting. [/INST]</s>"""
        
        # Add parameters for better text generation
        payload = {
            "inputs": formatted_prompt,
            "parameters": {
                "max_new_tokens": 500,
                "temperature": 0.7,
                "top_p": 0.95,
                "do_sample": True,
                "return_full_text": False
            }
        }
        
        response = requests.post(HF_API_URL, headers=headers, json=payload, timeout=30)
        
        if response.status_code != 200:
            logger.error(f"Hugging Face API returned status code {response.status_code}")
            logger.error(f"Response content: {response.text}")
            raise Exception(f"Hugging Face API error: {response.text}")
            
        response_data = response.json()
        logger.info(f"Received response from Hugging Face API: {str(response_data)[:200]}")
        
        # Handle response format
        if isinstance(response_data, list) and response_data:
            generated_text = response_data[0].get('generated_text', '')
            if not generated_text:
                raise Exception("Generated text is empty")
            
            # Clean up the generated text
            generated_text = generated_text.strip()
            # Remove any instruction markers that might have been generated
            generated_text = generated_text.replace('[INST]', '').replace('[/INST]', '')
            
            if not generated_text:
                raise Exception("Generated text is empty after cleanup")
            
            return generated_text
            
        raise Exception(f"Unexpected response format: {str(response_data)[:200]}")
    except requests.exceptions.Timeout:
        logger.error("Request to Hugging Face API timed out")
        raise Exception("Request timed out")
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error: {str(e)}")
        raise Exception(f"Request error: {str(e)}")
    except Exception as e:
        logger.error(f"Error calling Hugging Face API: {str(e)}")
        raise

def validate_input(data):
    """Validate the input data for resume generation."""
    required_fields = ['name', 'email', 'phone', 'location', 'summary', 'skills', 'experience']
    for field in required_fields:
        if field not in data:
            raise ValueError(f"Missing required field: {field}")
        if not data[field]:
            raise ValueError(f"Field cannot be empty: {field}")
    if not isinstance(data['skills'], list):
        raise ValueError("Skills must be a list")

def format_document(doc, data, content):
    """Format the document with proper resume styling."""
    try:
        # Add name as title
        title = doc.add_heading(data['name'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info section
        contact_info = doc.add_paragraph()
        contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info.add_run(f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('location', '')}")
        
        # Add summary section
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(data['summary'])
        
        # Add skills section
        doc.add_heading('Skills', level=1)
        skills_para = doc.add_paragraph(style='List Bullet')
        skills_para.add_run(', '.join(data['skills']))
        
        # Add experience section
        doc.add_heading('Professional Experience', level=1)
        for exp in data['experience']:
            doc.add_heading(f"{exp['title']} at {exp['company']}", level=2)
            doc.add_paragraph(f"Dates: {exp['dates']}")
            doc.add_paragraph(exp['description'])
        
        # Add generated content in a structured way
        sections = content.split('\n\n')
        for section in sections:
            if section.strip():
                doc.add_paragraph(section.strip())
        
        # Set consistent formatting
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
        
        logger.info("Document formatted successfully")
    except Exception as e:
        logger.error(f"Error formatting document: {str(e)}")
        raise

def get_style_config(style_data):
    color_schemes = {
        'classic': {'heading': RGBColor(0, 32, 96), 'subheading': RGBColor(0, 51, 153)},
        'modern': {'heading': RGBColor(0, 128, 128), 'subheading': RGBColor(0, 150, 150)},
        'professional': {'heading': RGBColor(128, 0, 0), 'subheading': RGBColor(153, 0, 0)},
        'elegant': {'heading': RGBColor(64, 64, 64), 'subheading': RGBColor(96, 96, 96)}
    }
    font_sizes = {
        'small': {'name': 16, 'heading': 12, 'normal': 10},
        'normal': {'name': 18, 'heading': 14, 'normal': 11},
        'large': {'name': 20, 'heading': 16, 'normal': 12}
    }
    return {
        'colors': color_schemes[style_data.get('colorScheme', 'classic')],
        'font': style_data.get('fontFamily', 'Calibri'),
        'sizes': font_sizes[style_data.get('fontSize', 'normal')]
    }

def enhance_content(section_type, content):
    """Enhance resume content using AI."""
    prompts = {
        'summary': f"As a professional resume writer, enhance this summary while keeping the same core information and being truthful: {content}",
        'experience': f"As a professional resume writer, enhance this job experience with strong action verbs and quantifiable achievements while keeping the same core information and being truthful: {content}",
        'skills': f"As a professional resume writer, organize and enhance this list of skills with industry-standard terminology while keeping the same core skills: {content}"
    }
    
    try:
        enhanced = generate_text(prompts[section_type])
        return enhanced if enhanced else content
    except Exception as e:
        logger.error(f"Error enhancing {section_type}: {str(e)}")
        return content

@app.route('/')
def index():
    return jsonify({"message": "AI Resume Builder API is running"})

@app.route('/test', methods=['GET'])
def test():
    return jsonify({"message": "API is working!"})

@app.route('/generate', methods=['POST'])
def generate_resume():
    try:
        data = request.get_json()
        app.logger.info(f"Received data: {data}")
        
        # Validate required fields
        required_fields = ['name', 'email', 'phone', 'location', 'summary', 'technicalSkills', 'softSkills', 'experience']
        if not all(field in data for field in required_fields):
            raise ValueError("Missing required fields")
            
        # Create document
        doc = Document()
        style = get_style_config(data.get('style', {}))
        
        # Set font for entire document
        style_font = doc.styles['Normal'].font
        style_font.name = style['font']
        
        # Name at the top
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(data['name'].upper())
        name_run.bold = True
        name_run.font.size = Pt(style['sizes']['name'])
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_info = f"{data['email']} | {data['phone']} | {data['location']}"
        contact_paragraph = doc.add_paragraph()
        contact_run = contact_paragraph.add_run(contact_info)
        contact_run.font.size = Pt(style['sizes']['normal'])
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing
        
        # Professional Summary
        summary_heading = doc.add_paragraph()
        summary_run = summary_heading.add_run('Professional Summary')
        summary_run.bold = True
        summary_run.font.color.rgb = style['colors']['heading']
        summary_run.font.size = Pt(style['sizes']['heading'])
        
        summary_paragraph = doc.add_paragraph()
        summary_run = summary_paragraph.add_run(data['summary'])
        summary_run.font.size = Pt(style['sizes']['normal'])
        
        # Technical & Soft Skills
        skills_heading = doc.add_paragraph()
        skills_run = skills_heading.add_run('Technical & Soft Skills')
        skills_run.bold = True
        skills_run.font.color.rgb = style['colors']['heading']
        skills_run.font.size = Pt(style['sizes']['heading'])
        
        # Technical Skills
        tech_skills_para = doc.add_paragraph()
        tech_skills_label = tech_skills_para.add_run('Technical Skills: ')
        tech_skills_label.bold = True
        tech_skills_para.add_run(', '.join(data['technicalSkills']))
        
        # Soft Skills
        soft_skills_para = doc.add_paragraph()
        soft_skills_label = soft_skills_para.add_run('Soft Skills: ')
        soft_skills_label.bold = True
        soft_skills_para.add_run(', '.join(data['softSkills']))
        
        # Education Section
        if 'education' in data and data['education']:
            education_heading = doc.add_paragraph()
            education_run = education_heading.add_run('Education')
            education_run.bold = True
            education_run.font.color.rgb = style['colors']['heading']
            education_run.font.size = Pt(style['sizes']['heading'])
            
            for edu in data['education']:
                edu_paragraph = doc.add_paragraph()
                school_run = edu_paragraph.add_run(edu['school'])
                school_run.bold = True
                school_run.font.size = Pt(style['sizes']['normal'])
                
                if edu['degree']:
                    edu_paragraph.add_run(' - ')
                    degree_run = edu_paragraph.add_run(edu['degree'])
                    degree_run.font.size = Pt(style['sizes']['normal'])
                
                if edu['graduationDate'] or edu.get('gpa'):
                    date_gpa = []
                    if edu['graduationDate']:
                        date_gpa.append(f"Graduation: {edu['graduationDate']}")
                    if edu.get('gpa'):
                        date_gpa.append(f"GPA: {edu['gpa']}")
                    
                    details_paragraph = doc.add_paragraph()
                    details_run = details_paragraph.add_run(' | '.join(date_gpa))
                    details_run.italic = True
                    details_run.font.size = Pt(style['sizes']['normal'])
        
        # Professional Experience
        experience_heading = doc.add_paragraph()
        experience_run = experience_heading.add_run('Professional Experience')
        experience_run.bold = True
        experience_run.font.color.rgb = style['colors']['heading']
        experience_run.font.size = Pt(style['sizes']['heading'])
        
        for exp in data['experience']:
            company_paragraph = doc.add_paragraph()
            company_run = company_paragraph.add_run(exp['company'])
            company_run.bold = True
            company_run.font.size = Pt(style['sizes']['normal'])
            
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(exp['title'])
            title_run.font.color.rgb = style['colors']['subheading']
            title_run.font.size = Pt(style['sizes']['normal'])
            
            if exp['dates']:
                dates_run = title_paragraph.add_run(f" | {exp['dates']}")
                dates_run.italic = True
                dates_run.font.size = Pt(style['sizes']['normal'])
            
            if exp['description']:
                desc_paragraph = doc.add_paragraph(style='List Bullet')
                desc_run = desc_paragraph.add_run(exp['description'])
                desc_run.font.size = Pt(style['sizes']['normal'])
        
        # Projects Section
        if data.get('projects'):
            projects_heading = doc.add_paragraph()
            projects_run = projects_heading.add_run('Projects')
            projects_run.bold = True
            projects_run.font.color.rgb = style['colors']['heading']
            projects_run.font.size = Pt(style['sizes']['heading'])
            
            for project in data['projects']:
                project_title = doc.add_paragraph()
                title_run = project_title.add_run(project['title'])
                title_run.bold = True
                title_run.font.size = Pt(style['sizes']['normal'])
                
                if project['technologies']:
                    tech_para = doc.add_paragraph()
                    tech_run = tech_para.add_run(f"Technologies: {project['technologies']}")
                    tech_run.italic = True
                    tech_run.font.size = Pt(style['sizes']['normal'])
                
                if project['description']:
                    desc_para = doc.add_paragraph(style='List Bullet')
                    desc_run = desc_para.add_run(project['description'])
                    desc_run.font.size = Pt(style['sizes']['normal'])
        
        # Certifications Section
        if data.get('certifications'):
            cert_heading = doc.add_paragraph()
            cert_run = cert_heading.add_run('Certifications / Online Courses')
            cert_run.bold = True
            cert_run.font.color.rgb = style['colors']['heading']
            cert_run.font.size = Pt(style['sizes']['heading'])
            
            for cert in data['certifications']:
                cert_para = doc.add_paragraph(style='List Bullet')
                name_run = cert_para.add_run(f"{cert['name']} - {cert['issuer']}")
                name_run.font.size = Pt(style['sizes']['normal'])
                if cert['date']:
                    date_run = cert_para.add_run(f" ({cert['date']})")
                    date_run.italic = True
                    date_run.font.size = Pt(style['sizes']['normal'])
        
        # Achievements Section
        if data.get('achievements'):
            achieve_heading = doc.add_paragraph()
            achieve_run = achieve_heading.add_run('Achievements / Awards')
            achieve_run.bold = True
            achieve_run.font.color.rgb = style['colors']['heading']
            achieve_run.font.size = Pt(style['sizes']['heading'])
            
            for achievement in data['achievements']:
                achieve_para = doc.add_paragraph(style='List Bullet')
                title_run = achieve_para.add_run(f"{achievement['title']}")
                title_run.bold = True
                title_run.font.size = Pt(style['sizes']['normal'])
                
                if achievement['date']:
                    date_run = achieve_para.add_run(f" ({achievement['date']})")
                    date_run.italic = True
                    date_run.font.size = Pt(style['sizes']['normal'])
                
                if achievement['description']:
                    desc_para = doc.add_paragraph()
                    desc_run = desc_para.add_run(achievement['description'])
                    desc_run.font.size = Pt(style['sizes']['normal'])
        
        # Activities Section
        if data.get('activities'):
            activities_heading = doc.add_paragraph()
            activities_run = activities_heading.add_run('Extra-curricular / Volunteer Activities')
            activities_run.bold = True
            activities_run.font.color.rgb = style['colors']['heading']
            activities_run.font.size = Pt(style['sizes']['heading'])
            
            for activity in data['activities']:
                activity_para = doc.add_paragraph(style='List Bullet')
                title_run = activity_para.add_run(f"{activity['title']} - {activity['organization']}")
                title_run.bold = True
                title_run.font.size = Pt(style['sizes']['normal'])
                
                if activity['date']:
                    date_run = activity_para.add_run(f" ({activity['date']})")
                    date_run.italic = True
                    date_run.font.size = Pt(style['sizes']['normal'])
                
                if activity['description']:
                    desc_para = doc.add_paragraph()
                    desc_run = desc_para.add_run(activity['description'])
                    desc_run.font.size = Pt(style['sizes']['normal'])
        
        # Languages Section
        if data.get('languages'):
            lang_heading = doc.add_paragraph()
            lang_run = lang_heading.add_run('Languages Known')
            lang_run.bold = True
            lang_run.font.color.rgb = style['colors']['heading']
            lang_run.font.size = Pt(style['sizes']['heading'])
            
            lang_para = doc.add_paragraph()
            for i, lang in enumerate(data['languages']):
                if i > 0:
                    lang_para.add_run(' | ')
                lang_para.add_run(f"{lang['name']} ({lang['proficiency']})")
        
        # Hobbies Section (Optional)
        if data.get('hobbies'):
            hobbies_heading = doc.add_paragraph()
            hobbies_run = hobbies_heading.add_run('Hobbies / Interests')
            hobbies_run.bold = True
            hobbies_run.font.color.rgb = style['colors']['heading']
            hobbies_run.font.size = Pt(style['sizes']['heading'])
            
            hobbies_para = doc.add_paragraph()
            hobbies_run = hobbies_para.add_run(data['hobbies'])
            hobbies_run.font.size = Pt(style['sizes']['normal'])
        
        # Save the document
        doc.save('resume.docx')
        app.logger.info("Document created successfully at " + os.path.abspath('resume.docx'))
        return jsonify({"message": "Resume generated successfully"})
        
    except Exception as e:
        app.logger.error(f"Error creating document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/download', methods=['GET'])
def download_resume():
    """Download the generated resume."""
    try:
        if not os.path.exists('resume.docx'):
            return jsonify({"error": "Resume file not found"}), 404
            
        directory = os.getcwd()
        return send_from_directory(
            directory,
            'resume.docx',
            as_attachment=True,
            download_name='resume.docx'
        )
    except Exception as e:
        logger.error(f"Error downloading resume: {str(e)}")
        return jsonify({"error": f"Failed to download resume: {str(e)}"}), 500

@app.after_request
def after_request(response):
    """Add CORS headers to all responses."""
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

if __name__ == '__main__':
    logger.info("Starting Flask server on http://localhost:5000")
    app.run(host='0.0.0.0', port=5000, debug=True) 